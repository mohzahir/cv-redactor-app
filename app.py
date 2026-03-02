import streamlit as st
import fitz  # PyMuPDF
import io
import zipfile
import pandas as pd
import json
import google.generativeai as genai
from docx import Document

st.set_page_config(page_title="AI CV Redactor & Extractor", page_icon="🤖")
st.title("AI-Powered Bulk CV Redactor")
st.write("Upload CVs to have Google Gemini intelligently find and redact contact info, while generating an Excel summary.")

# Safely load the API Key
try:
    genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
    model = genai.GenerativeModel('gemini-2.5-flash')
except Exception as e:
    st.error("⚠️ Gemini API Key not found. Please add it to Streamlit Secrets.")

uploaded_files = st.file_uploader("Upload candidate CVs", type=["pdf", "docx"], accept_multiple_files=True)

if uploaded_files:
    st.info(f"Processing {len(uploaded_files)} document(s)...")
    all_candidates_data = []

    try:
        zip_buffer = io.BytesIO()
        
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            for uploaded_file in uploaded_files:
                file_ext = uploaded_file.name.split('.')[-1].lower()
                output_buffer = io.BytesIO()
                output_filename = f"REDACTED_{uploaded_file.name}"
                
                doc_text = ""
                doc = None

                # ==========================================
                # PASS 1: EXTRACT ALL TEXT FOR THE AI
                # ==========================================
                file_bytes = uploaded_file.read()
                
                if file_ext == "pdf":
                    doc = fitz.open(stream=file_bytes, filetype="pdf")
                    for page in doc:
                        text = page.get_text("text")
                        if len(text.strip()) < 50:
                            try:
                                text = page.get_textpage_ocr(flags=0, language='eng', dpi=150, full=True).extractText()
                            except Exception:
                                pass
                        doc_text += text + "\n"
                        
                elif file_ext == "docx":
                    doc_docx = Document(io.BytesIO(file_bytes))
                    doc_text = "\n".join([para.text for para in doc_docx.paragraphs])

                # ==========================================
                # PASS 2: AI INTELLIGENCE & JSON EXTRACTION
                # ==========================================
                ai_prompt = f"""
                You are an expert recruitment assistant. Analyze the following CV text and extract the candidate's details into a strict JSON format.
                
                Use exactly these keys for the candidate data:
                "Name", "Qualification", "Age", "Email", "Phone", "Current Position", "Nationality", "Current Location".
                If a piece of information is missing, use "Not Found" as the value. Do not guess.

                CRITICAL INSTRUCTION FOR REDACTION:
                Create a 9th key named "Exact_Contacts_To_Redact". This must be a list of strings containing EVERY phone number, email address, and LinkedIn profile URL you found in the text.
                You MUST extract these strings EXACTLY as they appear in the text, character-for-character, including all spaces, dashes, or formatting. If the text says "+9 7 1 (50) 123", you must return exactly "+9 7 1 (50) 123". Do not fix or reformat them, or the system will fail to redact them.

                Return ONLY the raw JSON object, without markdown formatting.
                
                CV Text:
                {doc_text[:8000]}
                """
                
                try:
                    response = model.generate_content(ai_prompt)
                    json_text = response.text.strip().removeprefix('```json').removesuffix('```').strip()
                    extracted_data = json.loads(json_text)
                except Exception as e:
                    extracted_data = {
                        "Name": "AI Error", "Qualification": "AI Error", "Age": "AI Error", 
                        "Email": "AI Error", "Phone": "AI Error", "Current Position": "AI Error", 
                        "Nationality": "AI Error", "Current Location": "AI Error",
                        "Exact_Contacts_To_Redact": []
                    }

                strings_to_redact = extracted_data.get("Exact_Contacts_To_Redact", [])
                
                candidate_record = {"File Name": uploaded_file.name}
                for key in ["Name", "Qualification", "Age", "Email", "Phone", "Current Position", "Nationality", "Current Location"]:
                    candidate_record[key] = extracted_data.get(key, "Not Found")
                all_candidates_data.append(candidate_record)

                # ==========================================
                # PASS 3: REDACT THE EXACT STRINGS
                # ==========================================
                if file_ext == "pdf" and doc is not None:
                    for page in doc:
                        # Re-generate OCR textpage safely on the fly to avoid weak reference errors
                        tp = None
                        if len(page.get_text("text").strip()) < 50:
                            try:
                                tp = page.get_textpage_ocr(flags=0, language='eng', dpi=150, full=True)
                            except Exception:
                                pass
                        
                        for target_string in strings_to_redact:
                            if target_string and len(str(target_string).strip()) > 4:
                                text_instances = page.search_for(str(target_string), textpage=tp) if tp else page.search_for(str(target_string))
                                for inst in text_instances:
                                    page.add_redact_annot(inst, fill=(0, 0, 0))
                                    
                        page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_PIXELS)
                    doc.save(output_buffer, garbage=4, deflate=True)
                    doc.close()

                elif file_ext == "docx":
                    doc_docx = Document(io.BytesIO(file_bytes)) # Reload the Word doc safely
                    def replace_text_in_run(run):
                        for target_string in strings_to_redact:
                            if target_string and len(str(target_string).strip()) > 4:
                                if str(target_string) in run.text:
                                    run.text = run.text.replace(str(target_string), "[REDACTED]")

                    for para in doc_docx.paragraphs:
                        for run in para.runs:
                            replace_text_in_run(run)
                    for table in doc_docx.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        replace_text_in_run(run)
                    doc_docx.save(output_buffer)

                zip_file.writestr(output_filename, output_buffer.getvalue())

            # --- CREATE THE EXCEL FILE ---
            df = pd.DataFrame(all_candidates_data)
            excel_buffer = io.BytesIO()
            with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name='Candidates')
            zip_file.writestr("Candidate_Summary_Data.xlsx", excel_buffer.getvalue())

        st.success("All documents processed and redacted via AI!")
        st.download_button(
            label="Download Zip (Redacted CVs + AI Excel Data)",
            data=zip_buffer.getvalue(),
            file_name="AI_Processed_CVs.zip",
            mime="application/zip"
        )
        
    except Exception as e:
        st.error(f"An error occurred while processing: {e}")
